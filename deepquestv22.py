import streamlit as st
from dotenv import load_dotenv
from writer import report_writer
from planner import plan_research, replanner
from test import execute_step
from io import BytesIO
from docx import Document
from bs4 import BeautifulSoup
import markdown as md
import logging

load_dotenv()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s"
)

st.title("deepQuest v2")
st.sidebar.title("Research Steps")

def generate_word_doc_from_markdown(markdown_text):
    try:
        html = md.markdown(markdown_text, extensions=['tables'])
        soup = BeautifulSoup(html, "html.parser")
        doc = Document()
        doc.add_heading("DeepQuest Research Report", 0)

        for element in soup.children:
            if element.name and element.name.startswith("h") and element.name[1:].isdigit():
                level = int(element.name[1:])
                doc.add_heading(element.get_text(), level=level)
            elif element.name == "ul":
                for li in element.find_all("li"):
                    doc.add_paragraph(li.get_text(), style="List Bullet")
            elif element.name == "ol":
                for li in element.find_all("li"):
                    doc.add_paragraph(li.get_text(), style="List Number")
            elif element.name == "p":
                doc.add_paragraph(element.get_text())
            elif element.name == "table":
                rows = element.find_all("tr")
                if not rows:
                    continue
                cols = rows[0].find_all(["td", "th"])
                n_cols = len(cols)
                n_rows = len(rows)
                table = doc.add_table(rows=n_rows, cols=n_cols)
                for row_idx, row in enumerate(rows):
                    cells = row.find_all(["td", "th"])
                    for col_idx, cell in enumerate(cells):
                        table.cell(row_idx, col_idx).text = cell.get_text()
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        logging.error(f"Error converting markdown to Word: {e}")
        return None

# --- Session State Management ---
if "query" not in st.session_state:
    st.session_state.query = ""
if "steps" not in st.session_state:
    st.session_state.steps = []
if "completed_steps" not in st.session_state:
    st.session_state.completed_steps = []
if "context" not in st.session_state:
    st.session_state.context = ""
if "report" not in st.session_state:
    st.session_state.report = None

# query = st.chat_input("Enter your research query:")

# ...existing imports and setup...

query = st.chat_input("Enter your research query:")

# Set your max_steps dynamically or statically as needed
max_steps = 20  # Or use a value from Q-learning or user input

if not st.session_state.steps or st.session_state.query != query:
    st.session_state.steps = plan_research(query, max_steps=max_steps)
    st.session_state.completed_steps = []
    st.session_state.context = ""
    st.session_state.report = None

if query:
    st.session_state.query = query
    try:
        steps = st.session_state.steps
        sidebar_steps = st.sidebar.empty()
        sidebar_steps.markdown(
            "\n".join([f"{idx+1}. {step}" for idx, step in enumerate(steps)])
        )

        context = st.session_state.context
        completed_steps = st.session_state.completed_steps
        i = len(completed_steps)
        replan_rounds = 0
        replan_limit_reached = False
        max_steps_warning_shown = False

        progress_bar = st.progress(0, text="Starting research steps...")

        while i < len(steps):
            if len(steps) > max_steps and not max_steps_warning_shown:
                st.warning(
                    f"Maximum total steps ({max_steps}) reached. No further replanning will be done, but all planned steps will be executed."
                )
                max_steps_warning_shown = True
                replan_limit_reached = True

            step = steps[i]
            try:
                result = execute_step(step, context)
            except Exception as e:
                logging.error(f"Error executing step '{step}': {e}")
                st.error("Brain down, try again shortly!")
                st.stop()
            completed_steps.append((step, result))
            context += f"\nStep: {step}\nResult: {result}\n"

            # Update session state
            st.session_state.completed_steps = completed_steps
            st.session_state.context = context

            # Update plan display to show completed steps (with checkmark)
            plan_lines = []
            for idx, s in enumerate(steps):
                if idx < len(completed_steps):
                    plan_lines.append(f"✅ **Step {idx+1}:** {s}\n")
                else:
                    plan_lines.append(f"**Step {idx+1}:** {s}\n")
            sidebar_steps.markdown(
                "\n".join(
                    [
                        (
                            f"✅ {idx+1}. {s}\n"
                            if idx < len(completed_steps)
                            else f"{idx+1}. {s}"
                        )
                        for idx, s in enumerate(steps)
                    ]
                )
            )

            # Update progress bar
            progress = int((len(completed_steps) / len(steps)) * 100)
            progress_bar.progress(progress / 100, text=f"Completed {len(completed_steps)} of {len(steps)} steps")

            # Replanning
            if not replan_limit_reached:
                try:
                    steps, replan_rounds, replan_limit_reached = replanner(
                        context, steps, replan_rounds, 3, replan_limit_reached, max_steps=max_steps
                    )
                    st.session_state.steps = steps
                except Exception as e:
                    logging.error(f"Error during replanning: {e}")
                    st.error("Brain down, try again shortly!")
                    st.stop()
            i += 1

        progress_bar.progress(1.0, text="All steps completed!")

        # Generate report only if not already in session state
        if not st.session_state.report:
            try:
                st.session_state.report = report_writer(context)
            except Exception as e:
                logging.error(f"Error generating report: {e}")
                st.error("Brain down, try again shortly!")
                st.stop()

    except Exception as e:
        logging.critical(f"Critical error in main UI: {e}")
        st.error("Brain down, try again shortly!")

# --- Always display report and download button if available ---
if st.session_state.report:
    st.subheader("Final Research Report")
    st.markdown(st.session_state.report)
    word_buffer = generate_word_doc_from_markdown(st.session_state.report)
    if word_buffer:
        st.download_button(
            label="Download Report as Word Document",
            data=word_buffer,
            file_name="deepquest_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        st.error("Brain down, try again shortly!")